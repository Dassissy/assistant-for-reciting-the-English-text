# -*- coding: utf-8 -*-
"""
Created on Sun Apr 18 20:15:38 2021

英语默写小助手
问题还有：绿色太多（设空句的其它内容变成绿色）、标点更换后大小写出问题、破折号的处理不能这么简单、文件未关闭要处理一下、
         配置文件压根没设置、没了标点没内味
"""
from docx import Document
from docx.shared import RGBColor
import random,re,time,os
  
def change_words(s):#更改标点，但我不会拼
    try:
        s = re.sub("!",".",s)#把叹号换成句号
    except:
        pass
    try:
        s = re.sub("?",".",s)#问号也换掉
    except:
        pass
    try:
        s = re.sub("--"," ",s)#破折号换空格
    except:
        pass
    try:
        s = re.sub("\""," ",s)#引号换空格
    except:
        pass
    try:
        s = re.sub(","," ",s)#逗号换空格
    except:
        pass  
    try:
        s = re.sub(r"[ ]{2,}"," ",s)#多于一个的空格换成单个空格
    except:
        pass
    return s    
def get_s(text_path,FROM=2,TO=5):
    document = Document(text_path)
    s = ''
    for para in document.paragraphs:#读取word文档
        s = s + para.text
    s = change_words(s)#更改标点
    sens = s.split(".")#按句分割
    new_sens = []
    for sen in sens:
        if not sen == ' ' or sen == '':
            new_sens.append(sen)
    #获取一些幸运句子
    #len返回的值是从1开始数的，随机数是含末尾的（len(new_sens)-1)
    lucky_list = [random.randint(0, len(new_sens)-1) for _ in range(len(new_sens)//2)]
    SET = set(lucky_list)
    lucky_list = list(SET)#利用集合去重
    lucky_list.sort()#排序
    luck_len_dict = {}
    no_list = []#某些句子长度只有1
    answer = new_sens.copy()
    print("answer is:{}".format(answer))
    lucky_word_list = []#记录lucky_word
    for i in lucky_list:
        print(("i and lucky_list are:{},{}").format(i,lucky_list))
        sen = new_sens[i]
        luck_len = random.randint(FROM,TO)#挖空长度
        word_list = sen.split(" ")
        LEN = len(word_list)
        S_len = len(sen.split(" "))
        while luck_len > 1 and luck_len > S_len:
            luck_len -= 1#确保挖的空不会超出句子
        if LEN >= FROM:#若到最后LEN仍然在FROM以上
            lucky_word = random.randint(0, LEN-luck_len)#index要少一个，但len()会多一个
            lucky_word_list.append(lucky_word)
        else:
            no_list.append(i)
            continue
        luck_len_dict[i] = luck_len
        word_list[lucky_word] = "({})".format(luck_len)#括号中间写长度
        for i2 in range(luck_len):
            if not i2 == 0:#不能再把之前的括号删掉
                word_list[lucky_word+i2] = ''#删掉之后的
        sen = ''
        for word in word_list:
            sen = sen + word + " "
        new_sens[i] = sen#处理后的句子再转回去
    for no in no_list:
        lucky_list.append(no)
    s = ''
    for sen in new_sens:
        s = s + sen + "."
    s = re.sub(r"[ ]{2,}"," ",s)#多于一个的空格转为一个空格
    return s,answer,lucky_list,luck_len_dict,lucky_word_list

def dig_hole(text_path,file_path):#挖空（洞）
    '''while True:
        try:#从根本上解决出错的问题
            s,answer,lucky_list,luck_len_dict,lucky_word_list = get_s(text_path)
            break
        except:
            continue#多跑几遍绝对没问题(doge)'''
    s,answer,lucky_list,luck_len_dict,lucky_word_list = get_s(text_path)
    lucky_list = set(lucky_list)
    lucky_list = list(lucky_list)#不知为何这里仍有重复
    print("lucky_list after duplicate removal is:{}".format(lucky_list))
    path = file_path
    document = Document()
    document.save(path)
    document = Document(path)#再开一遍，清空内容
    document.add_paragraph(s)
    document.save(path)
    return answer,lucky_list,luck_len_dict,lucky_word_list,s

def correct(answer,lucky_list,luck_len_dict,lucky_word_list,s,file_path):
    document = Document(file_path)
    s = ''
    for para in document.paragraphs:
        s = s + para.text
    s = change_words(s)#更改标点
    s_list = s.split(".")#同样按句分割答案
    for i in range(len(s_list)):#最后会多一个空格
        list_s = list(s_list[i])
        if len(list_s) >= 1:
            while list_s[-1] == " ":
                list_s.pop()
        _s_ = ''
        for w in list_s:
            _s_ = _s_ + w
        s_list[i] = _s_
    er = 0#error
    al = 0#all
    for key in luck_len_dict:
        al += luck_len_dict[key]#不要+1！！！
    error_dict = {}#对了用绿色,错了用红色，精确到单词
    correct_list = []#正确的句子用绿色
    ilist = []
    for i in range(len(lucky_list)):
        if lucky_list[i] not in luck_len_dict:#弹出放弃的句子（过短）
            ilist.append(i)
    print(("ilist is:{}\nlucky_list is:{}\nluck_len_dict is:{}").format(ilist,lucky_list,luck_len_dict))
    count = 0#使用count抵消掉弹出后列表长度的缩减，这里又可以用了
    for i in ilist:
        lucky_list.pop(i-count)
        count += 1
    for i in range(len(lucky_list)):
        lucky_list[i] = (lucky_list[i],i)
    print(("lucky_list with count is:{}").format(lucky_list))
    """
    曾使用额外初始化一个count的方法，但是发生了错误，上面又没出错，现象：列表末位离奇+1,（0,1,2,4）,返回如下：
    3 5
    0 4
    1 5
    [3, 0, 1, 0] {1: 5, 8: 4, 11: 5, 15: 4} 15 4
    测试用代码：
    try:
        print(lucky_word_list[count],luck_len_dict[i])
    except:
        print(lucky_word_list,luck_len_dict,i,count)
    """
    for i,count in lucky_list:
        if s_list[i] == answer[i]:#如果对了就过了
            correct_list.append(i)
            continue
        else:#没对就费事了
            error_dict[i] = []#往字典里加空列表
            error_words = s_list[i].split(" ")
            correct_words = answer[i].split(" ")
            if len(error_words) == len(correct_words):#长度相等
                for j in range(len(error_words)):
                    if not error_words[j] == correct_words[j]:#没对的话：)
                        print(("error_words[j] and correct_words[i] are:{},{}").format(error_words[j],correct_words[j]))
                        error_dict[i].append(j)#记下应标红的单词，及其所在的句子
            elif len(error_words) != len(correct_words):#不一样的话：
                A = 1
                try:
                    print(("lucky_word_list[count],luck_len_dict[i] are:{},{}").foramt(lucky_word_list[count],luck_len_dict[i]))
                except:
                    print(("lucky_word_list,luck_len_dict,i,count are:{},{},{},{}").format(lucky_word_list,luck_len_dict,i,count))
                    pass
                for j in range(lucky_word_list[count],lucky_word_list[count]+luck_len_dict[i]):
                    a = (lucky_word_list[count],lucky_word_list[count]+luck_len_dict[i])
                    error_dict[i].append(j)
            '''这部分没有用了，哈哈哈
            else:#填多了（为什么会填多了？）
                print(error_words,correct_words)
                f=0
                b=0
                for j in range(len(error_words)):#从前面走一遍
                    if error_words[j] == correct_words[j]:
                        continue
                    else:
                        f = j
                        break
                error_words.reverse()#倒转原答
                correct_words.reverse()#倒转标答
                sdrow_rorre = error_words.copy()
                sdrow_tcerroc = correct_words.copy()
                error_words.reverse()#倒转回原答
                correct_words.reverse()#倒转回标答
                for j in range(len(sdrow_rorre)):#从后面走一遍
                    if sdrow_rorre[j] == sdrow_tcerroc[j]:
                        continue
                    else:
                        b = j
                        break
                b = len(error_words) - b
                print(f,b)
                if f != b:
                    for j in range(f,b+1):
                        print(error_dict[i])
                        error_dict[i].append(j)
                else:
                    error_dict[i].append(f)
            '''
    document = Document()
    document.save(file_path)
    document = Document(file_path)#再开一遍，清空内容
    p = document.add_paragraph()
    while s_list[-1] == "":
        s_list.pop()
    for i in range(len(lucky_list)):
        lucky_list[i] = lucky_list[i][0]#又要从元组变回单个的数字
    for i in range(len(answer)):
        print(("i and lukcy_list are:{} and {}:").format(i,lucky_list))
        if i not in lucky_list:
            print(("i:{} is not in lucky_list").format(i))
            try:
                print(("answer[i] is:{}").format(answer[i]))
            except:
                print(("s_list[i] is:{}").format(s_list[i]))
            r = p.add_run(answer[i]+".")#没有挖空的直接放进去就好了，记得补句点
        elif i in correct_list:
            b = "sign"
            print(("i,correct_list are:{},{}").format(i,correct_list))
            r = p.add_run(answer[i]+".")
            r.font.color.rgb = RGBColor(0,255,0)#写对的句子用绿色！
        else:
            
            for I in range(len(answer[i].split(" "))):
                print(("i and error_dict are:{} and\n{}").format(i,error_dict))
                print(("error_dict[i] is:{}").format(error_dict[i]))
                if I == len(answer[i].split(" "))-1:
                    r = p.add_run(answer[i].split(" ")[I]+".")#补句点
                else:
                    r = p.add_run(answer[i].split(" ")[I]+" ")#补空格
                if not I in error_dict[i]:#往字典里加列表为的就是这个
                    r.font.color.rgb = RGBColor(0,255,0)#正确单词用绿色
                else:
                    r.font.color.rgb = RGBColor(255,0,0)#错误单词用红色
                    er += 1
    document.save(file_path)
    if (al-er)/al == 0:
        return 0
    else:
        wao = str(answer)+"\n"+str(lucky_list)+"\n"+str(error_dict)+"\n"+str(error_words)
        print("they are:answer, lucky_list,error_dict,error_words\n{}".format(wao))
        return (al-er)/al #返回正确率
def out(accuracy):
    if accuracy == 1:
        print("全对，真棒！")
    elif accuracy == 0:
        print("全错，真有你的")
    else:
        print("改完了，{}分，去看看吧".format(accuracy*100))

def main(text_path="D://英语课文//原文//",file_path="D://英语课文//答卷//"):
    path_list = []
    whole_path_list = []
    for path in os.listdir(text_path):
        T_path = text_path + path
        F_path = file_path + path
        path_list.append(path)
        whole_path_list.append((T_path,F_path))
    while True:
        try:
            for path in path_list:
                print(path)
            index = int(input("请选择文章（序号）：")) - 1#从一开始排序，所以减一
            if index > len(path_list)-1:#len()返回的值比索引大1（重要的东西写三遍）
                path_list.error()#只是出个错而已
            break
        except:
            print("重写")
    T_path,F_path = whole_path_list[index]
    """
    while True:
        try:
            answer,lucky_list,luck_len_dict,lucky_word,s = dig_hole(text_path=T_path,file_path=F_path)
            break
        except:
            print("\r文件没有关",end='')
            time.sleep(0.5)
            """
    answer,lucky_list,luck_len_dict,lucky_word,s = dig_hole(text_path=T_path,file_path=F_path)
    while True:
        try:
            accuracy = input("开始做吧,做好了就随便打些东西进来：")#用后面出现的变量是为了不报错且不影响
            accuracy = correct(answer,lucky_list,luck_len_dict,lucky_word,s,file_path=F_path)
            break
        except PermissionError:
            print("文件没有关")
    out(accuracy)
    
yes = ''
while yes == '':
    main()
    yes = input("还继续吗？继续就回车，否则按任意键后回车")